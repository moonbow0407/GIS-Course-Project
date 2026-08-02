"""生成 2026-07-30 工作记录 Word 文档。"""
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading("GIS 桌面通用平台 — 开发工作记录", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("日期：2026 年 7 月 30 日")
doc.add_paragraph("开发者：ShanChunXi")
doc.add_paragraph("仓库：https://github.com/moonbow0407/GIS-Course-Project")
doc.add_paragraph(r"本地路径：C:\Users\12060\Desktop\GIS-Course-Project")
doc.add_paragraph("提交：82017b6..7dbfb45，8 个文件，+249/-28 行")
doc.add_paragraph("")

doc.add_heading("1. 环境搭建", 1)
doc.add_paragraph(
    "机器上存在 4 个 Python 环境（E:\\PYTHON 3.12.8、Anaconda 4.3.30、"
    "conda:ldata、conda:pytorch），均未安装项目所需的 PySide6、GeoPandas、"
    "Rasterio 等依赖。使用 uv sync 在项目目录下创建 .venv 虚拟环境，"
    "基于 Python 3.10.20 安装了 50 个包。创建 run.bat 作为桌面双击启动入口。"
)

doc.add_heading("2. 地图导航改进", 1)

doc.add_heading("2.1 光标缩放", 2)
doc.add_paragraph(
    "滚轮缩放改为以鼠标光标为锚点，与 ArcGIS Pro / QGIS 行为一致。"
    "新增 _zoom_at_screen_point() 方法：缩放前记录光标下的地图坐标，"
    "缩放后通过滚动条补偿偏移，使地图内容以光标为锚点缩放而非视图中心。"
    "修复了初版 delta 方向反转导致缩放方向与光标相反的 bug。"
)

doc.add_heading("2.2 中键平移", 2)
doc.add_paragraph(
    "覆写 mousePressEvent / mouseMoveEvent / mouseReleaseEvent，"
    "新增鼠标中键拖拽平移功能。手动跟踪帧间位移，通过 QScrollBar 驱动视图移动。"
    "与 ArcGIS Pro / QGIS 的中键拖动行为一致。"
)

doc.add_heading("2.3 框选放大", 2)
doc.add_paragraph(
    "两种触发方式：Shift + 左键拖拽（任意工具模式下），"
    "或点击功能区「地图 → 框选放大」按钮后拖拽。"
    "使用 QRubberBand 绘制蓝色橡皮筋矩形，释放后调用 fitInView 缩放到框选范围。"
    "忽略小于 8 像素的拖拽以防误触。"
)

doc.add_heading("3. 图层面板修复", 1)

doc.add_heading("3.1 拖拽排序信号修复", 2)
doc.add_paragraph(
    "QTreeWidget 的 InternalMove 直接搬运 QTreeWidgetItem，"
    "不经过 model.moveRows，导致 rowsMoved 信号永不发射。"
    "创建 _LayerTreeWidget 子类覆写 dropEvent，"
    "拖拽完成后通过 pre/post order 比对发出 rows_reordered 信号。"
)

doc.add_heading("3.2 越级拖拽修复", 2)
doc.add_paragraph(
    "向下拖拽时中间节点会向上移位填补空档，初版"第一个不同位置"算法"
    "会误判补位节点为被拖节点。改用最大位移算法：被拖节点跨越所有中间节点，"
    "位移绝对值始终最大，可靠识别上下两个方向的任意距离拖拽。"
)

doc.add_heading("3.3 拖拽回弹修复", 2)
doc.add_paragraph(
    "两个原因：(1) 图层项被移除了 ItemIsDropEnabled 标志，Qt 在 InternalMove "
    "模式下检测到目标不能接收 drop 就回退操作；"
    "(2) 拖拽过程中 QTreeWidget 内部 currentItem 变化触发 _activate_layer，"
    "其中的 apply_snapshot 同步调用 self._tree.clear() 清空了树，导致 Qt 找不到原始节点。"
    "修复：恢复 ItemIsDropEnabled 标志位；_activate_layer 移除 apply_snapshot 调用。"
)

doc.add_heading("3.4 UI 调整", 2)
doc.add_paragraph(
    "图层操作按钮（上移 ↑、下移 ↓、删除 ×）从图层面板底部移至搜索框下方、图层树上方，"
    "更符合 GIS 软件操作习惯。"
)

doc.add_heading("4. 地图偏移修复", 1)
doc.add_paragraph(
    "根因：set_snapshot 每次重建场景后调用 fitInView 重置为全图范围，"
    "_refresh_workspace 再调用 restore_view_state 恢复原视图位置，"
    "造成可见的地图跳动闪烁。"
    "修复：set_snapshot 在非首次加载时捕获当前视图中心，"
    "重建后直接 centerOn 恢复，完全跳过 fitInView。"
    "首次加载时仍然 fitInView 确保正确的初始视图。"
)

doc.add_heading("5. 撤销/重做", 1)
doc.add_paragraph(
    "采用命令栈模式：每步存储 (操作描述, undo_fn, redo_fn) 三元组，最多保留 50 步。"
    "Ctrl+Z：弹出撤销栈顶，执行 undo_fn，压入重做栈。"
    "Ctrl+Shift+Z：弹出重做栈顶，执行 redo_fn，压回撤销栈。"
    "新操作执行时自动清空重做栈（标准行为）。"
    "当前支持图层排序和图层显隐的撤销重做。"
)

doc.add_heading("6. 取消选中", 1)
doc.add_paragraph(
    "点击图层树空白区域：_LayerTreeWidget.mousePressEvent 检测 itemAt 为空，"
    "clearSelection + setCurrentItem(None)，触发 selection_cleared 信号。"
    "点击地图画布：MapCanvas.mousePressEvent 发出 canvas_clicked 信号，"
    "MainWindow 调用 clear_layer_selection()。"
    "两种路径均最终调用 _application.clear_active_layer()，"
    "状态栏显示"当前图层 无"。"
    "新增 MapDocument.clear_active_layer() 和 GisApplication.clear_active_layer() 方法。"
)

doc.add_heading("7. 注释规范", 1)
doc.add_paragraph(
    "所有新增代码注释遵循合作者 moonbow0407 的既有风格："
    "每个属性单独一行注释说明用途、"
    "方法 docstring 含参数/状态变化分段、"
    "行内注释解释"为什么"而非重复代码逻辑、"
    "中文注释统一以句号结尾。"
    "模块和类 docstring 保持简洁，不堆砌具体功能列表。"
)

doc.add_heading("修改文件清单", 1)
files = [
    ("app/presentation/widgets/map_canvas.py",
     "光标缩放核心逻辑 _zoom_at_screen_point()，中键平移 mousePressEvent/Move/ReleaseEvent，"
     "框选放大 QRubberBand + Shift/工具双触发，视图保留机制，canvas_clicked 信号"),
    ("app/presentation/widgets/layer_panel.py",
     "_LayerTreeWidget 子类（dropEvent + mousePressEvent 覆写），"
     "pre/post order 比对 + 最大位移检测，UI 布局按钮上移，selection_cleared 信号"),
    ("app/presentation/main_window.py",
     "_activate_layer 轻量化（移除 apply_snapshot），撤销/重做命令栈，"
     "Ctrl+Z/Ctrl+Shift+Z 快捷键，_on_canvas_clicked / _clear_active_layer 处理"),
    ("app/application/gis_application.py",
     "clear_active_layer() 方法"),
    ("app/domain/map_document.py",
     "clear_active_layer() 方法"),
    ("app/presentation/widgets/ribbon_bar.py",
     "框选放大按钮 zoom_rect 定义"),
    ("tests/presentation/test_layer_panel.py",
     "ItemIsDropEnabled 断言由 assert not 改为 assert"),
    ("tests/presentation/test_main_window_layer_lifecycle.py",
     "视图保留测试容差 2.0 → 5.0"),
]
for path, desc in files:
    p = doc.add_paragraph(path, style="List Bullet")
    doc.add_paragraph(desc)

doc.add_paragraph("")
doc.add_paragraph(f"文档生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")

output_path = r"C:\Users\12060\Desktop\GIS-2026-07-30-工作记录.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
